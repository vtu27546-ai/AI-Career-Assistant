import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
import os
from PyPDF2 import PdfReader
from docx import Document
from reportlab.pdfgen import canvas
import pyttsx3
import threading

from auth import login_register_ui

if "logged_in" not in st.session_state:
    st.session_state.logged_in=False

if not st.session_state.logged_in:
    login_register_ui()
    st.stop()

# =========================================================
# LOAD ENVIRONMENT VARIABLES
# =========================================================
load_dotenv()

client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY")
)

# =========================================================
# SMOOTH FEMALE VOICE ASSISTANT
# =========================================================

def speak(text):

    def run_voice():

        try:

            engine = pyttsx3.init()

            voices = engine.getProperty('voices')

            # Female Voice
            if len(voices) > 1:
                engine.setProperty('voice', voices[1].id)

            # Smooth Voice Settings
            engine.setProperty('rate', 155)

            engine.setProperty('volume', 0.9)

            engine.say(text)

            engine.runAndWait()

            engine.stop()

        except Exception as e:

            print(e)

    threading.Thread(target=run_voice).start()

# =========================================================
# PAGE CONFIG
# =========================================================
st.set_page_config(
    page_title="AI Career Assistant Platform",
    page_icon="🚀",
    layout="wide"
)

# =========================================================
# CUSTOM CSS
# =========================================================
st.markdown("""
<style>

.main {
    background-color: #0E1117;
    color: white;
}

.stButton>button {
    background-color: #00ADB5;
    color: white;
    border-radius: 12px;
    height: 50px;
    width: 250px;
    font-size: 18px;
    border: none;
    font-weight: bold;
}

.skill-box {
    background-color: #1F2937;
    padding: 15px;
    border-radius: 12px;
    margin: 5px;
    font-size: 18px;
    color: white;
}

</style>
""", unsafe_allow_html=True)

# =========================================================
# TITLE
# =========================================================
st.title("🚀 AI Career Assistant Platform")

st.success(
    "ATS Analyzer + Resume Builder + AI Interview Coach + AI Career Mentor"
)

with st.sidebar:
    st.success(f"Logged in as {st.session_state.get('user_name','User')}")
    if st.button("🚪 Logout"):
        st.session_state.logged_in=False
        st.session_state.pop("user_name",None)
        st.rerun()


# =========================================================
# WELCOME VOICE
# =========================================================
if "welcome_done" not in st.session_state:

    speak(
        "Welcome to the AI Career Assistant Platform"
    )

    st.session_state.welcome_done = True

# =========================================================
# FILE UPLOAD
# =========================================================
uploaded_file = st.file_uploader(
    "📄 Upload Resume",
    type=["pdf", "docx"]
)

# =========================================================
# JOB DESCRIPTION
# =========================================================
job_description = st.text_area(
    "📋 Paste Job Description",
    height=250
)

# =========================================================
# EXTRACT RESUME TEXT
# =========================================================
resume_text = ""

if uploaded_file:

    # PDF
    if uploaded_file.name.endswith(".pdf"):

        pdf_reader = PdfReader(uploaded_file)

        for page in pdf_reader.pages:

            text = page.extract_text()

            if text:
                resume_text += text

    # DOCX
    elif uploaded_file.name.endswith(".docx"):

        doc = Document(uploaded_file)

        for para in doc.paragraphs:
            resume_text += para.text + "\n"

# =========================================================
# USER NAME
# =========================================================
user_name = "Candidate"

try:

    first_line = resume_text.split("\n")[0]

    if len(first_line) < 30:
        user_name = first_line

except:
    pass

# =========================================================
# USER GREETING
# =========================================================
if uploaded_file:

    st.success(f"👋 Welcome {user_name}")

# =========================================================
# TABS
# =========================================================
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 ATS Analysis",
    "📄 Modified Resume",
    "🎤 Interview Questions",
    "🤖 Personal AI Assistant",
    "📄 Cover Letter Generator"
])

# =========================================================
# ATS ANALYSIS
# =========================================================
with tab1:

    if st.button("Analyze Resume"):

        if uploaded_file and job_description:

            try:

                with st.spinner("Analyzing Resume..."):

                    skills_database = [
                        "python",
                        "sql",
                        "machine learning",
                        "data analysis",
                        "communication",
                        "problem solving",
                        "docker",
                        "aws",
                        "kubernetes",
                        "flask",
                        "django",
                        "react",
                        "streamlit",
                        "openai",
                        "api",
                        "html",
                        "css",
                        "javascript",
                        "git",
                        "github",
                        "pandas",
                        "numpy",
                        "tensorflow",
                        "deep learning",
                        "nlp",
                        "power bi",
                        "excel",
                        "java",
                        "c++",
                        "mongodb"
                    ]

                    resume_lower = resume_text.lower()
                    jd_lower = job_description.lower()

                    matching_skills = []
                    missing_skills = []

                    for skill in skills_database:

                        if skill in resume_lower and skill in jd_lower:
                            matching_skills.append(skill.title())

                        elif skill in jd_lower and skill not in resume_lower:
                            missing_skills.append(skill.title())

                    total_required = len(matching_skills) + len(missing_skills)

                    if total_required > 0:

                        ats_score = int(
                            (len(matching_skills) / total_required) * 100
                        )

                    else:

                        ats_score = 50

                    prompt = f"""
                    You are an expert ATS Resume Analyzer,
                    Recruiter, Career Coach, and HR Manager.

                    Analyze the resume deeply and professionally.

                    Give:
                    1. ATS Improvement Suggestions
                    2. Missing Important Keywords
                    3. Resume Strengths
                    4. Weak Areas
                    5. Recruiter Perspective
                    6. Career Advice
                    7. Industry-Level Improvements

                    Resume:
                    {resume_text}

                    Job Description:
                    {job_description}
                    """

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "You are a professional ATS resume expert and recruiter."
                            },
                            {
                                "role": "user",
                                "content": prompt
                            }
                        ]
                    )

                    result = response.choices[0].message.content

                st.subheader("📊 ATS Score")

                st.progress(ats_score / 100)

                if ats_score >= 80:

                    st.success(f"Excellent ATS Score: {ats_score}%")

                elif ats_score >= 60:

                    st.warning(f"Good ATS Score: {ats_score}%")

                else:

                    st.error(f"Low ATS Score: {ats_score}%")

                # MATCHING SKILLS
                st.subheader("✅ Matching Skills")

                if matching_skills:

                    cols = st.columns(2)

                    for i, skill in enumerate(matching_skills):

                        cols[i % 2].markdown(
                            f"<div class='skill-box'>✅ {skill}</div>",
                            unsafe_allow_html=True
                        )

                else:

                    st.warning("No matching skills found")

                # MISSING SKILLS
                st.subheader("❌ Missing Skills")

                if missing_skills:

                    cols2 = st.columns(2)

                    for i, skill in enumerate(missing_skills):

                        cols2[i % 2].markdown(
                            f"<div class='skill-box'>❌ {skill}</div>",
                            unsafe_allow_html=True
                        )

                else:

                    st.success("No missing skills")

                # AI Suggestions
                st.subheader("📋 AI Suggestions")

                st.write(result)

                # Confidence Booster
                st.subheader("🔥 Confidence Booster")

                confidence_lines = [
                    "You are skilled and capable",
                    "Believe in yourself",
                    "Confidence is the key to success",
                    "You can crack this interview",
                    "Stay calm and answer confidently"
                ]

                for line in confidence_lines:
                    st.success(line)

                speak(
                    f"{user_name}, you are doing great. Believe in yourself. You can crack this interview."
                )

            except Exception as e:

                st.error(f"Error: {e}")

        else:

            st.error(
                "Please upload resume and paste job description."
            )

# =========================================================
# MODIFIED RESUME
# =========================================================
with tab2:

    if st.button("Generate Modified Resume"):

        if uploaded_file and job_description:

            try:

                with st.spinner("Generating Modified Resume..."):

                    rewrite_prompt = f"""
                    You are an expert ATS Resume Writer.

                    Rewrite the resume professionally.

                    Requirements:
                    - Improve ATS compatibility
                    - Use strong action verbs
                    - Add recruiter-friendly wording
                    - Improve professionalism
                    - Keep proper formatting
                    - Optimize for this job description

                    Resume:
                    {resume_text}

                    Job Description:
                    {job_description}
                    """

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "You are a professional resume writer and recruiter."
                            },
                            {
                                "role": "user",
                                "content": rewrite_prompt
                            }
                        ]
                    )

                    modified_resume = response.choices[0].message.content

                st.subheader("📄 Modified Resume")

                st.write(modified_resume)

                # SAVE DOCX
                doc = Document()

                doc.add_paragraph(modified_resume)

                doc.save("Modified_Resume.docx")

                # SAVE PDF
                pdf_file = "Modified_Resume.pdf"

                c = canvas.Canvas(pdf_file)

                textobject = c.beginText(40, 800)

                for line in modified_resume.split('\n'):
                    textobject.textLine(line)

                c.drawText(textobject)

                c.save()

                # DOWNLOAD DOCX
                with open("Modified_Resume.docx", "rb") as file:

                    st.download_button(
                        label="⬇ Download DOCX Resume",
                        data=file,
                        file_name="Modified_Resume.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # DOWNLOAD PDF
                with open("Modified_Resume.pdf", "rb") as file:

                    st.download_button(
                        label="⬇ Download PDF Resume",
                        data=file,
                        file_name="Modified_Resume.pdf",
                        mime="application/pdf"
                    )

                st.success(
                    "🎉 Resume Generated Successfully"
                )

                speak(
                    f"Thank you {user_name}. Your modified resume is ready."
                )

            except Exception as e:

                st.error(f"Error: {e}")

        else:

            st.error(
                "Please upload resume and paste job description."
            )

# =========================================================
# INTERVIEW QUESTIONS
# =========================================================
with tab3:

    if st.button("Generate Interview Questions"):

        if uploaded_file and job_description:

            try:

                with st.spinner(
                    "Generating Interview Questions..."
                ):

                    interview_prompt = f"""
                    You are a senior technical interviewer and HR recruiter.

                    Generate:
                    1. 5 HR Questions with detailed answers
                    2. 5 Technical Questions with detailed answers
                    3. Confidence Tips
                    4. Communication Tips
                    5. Common Mistakes to Avoid

                    Resume:
                    {resume_text}

                    Job Description:
                    {job_description}
                    """

                    interview_response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "You are an expert interviewer and career coach."
                            },
                            {
                                "role": "user",
                                "content": interview_prompt
                            }
                        ]
                    )

                    questions = interview_response.choices[0].message.content

                st.subheader(
                    "🎤 Interview Questions & Answers"
                )

                st.write(questions)

                speak(
                    f"{user_name}, practice confidently. You will perform very well in your interview."
                )

            except Exception as e:

                st.error(f"Error: {e}")

        else:

            st.error(
                "Please upload resume and paste job description."
            )

# =========================================================
# PERSONAL AI ASSISTANT
# =========================================================
with tab4:

    st.subheader("🤖 Personal AI Career Assistant")

    st.write(
        "Ask anything about careers, coding, interviews, resumes, internships, and projects."
    )

    if "messages" not in st.session_state:
        st.session_state.messages = []

    for message in st.session_state.messages:

        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    user_input = st.chat_input(
        "Ask your AI Assistant..."
    )

    if user_input:

        st.session_state.messages.append(
            {
                "role": "user",
                "content": user_input
            }
        )

        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):

            with st.spinner("AI Assistant Thinking..."):

                assistant_prompt = f"""
                You are an advanced AI Career Mentor,
                Resume Expert,
                Interview Coach,
                Coding Guide,
                and Professional Recruiter.

                User Resume:
                {resume_text}

                User Question:
                {user_input}
                """

                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are a smart AI career mentor."
                        },
                        {
                            "role": "user",
                            "content": assistant_prompt
                        }
                    ]
                )

                assistant_reply = response.choices[0].message.content

                st.markdown(assistant_reply)

                st.session_state.messages.append(
                    {
                        "role": "assistant",
                        "content": assistant_reply
                    }
                )

                short_reply = assistant_reply[:200]

                speak(short_reply)

# =========================================================
# COVER LETTER GENERATOR
# =========================================================
with tab5:

    st.subheader("📄 AI Cover Letter Generator")

    st.write(
        "Generate professional AI-powered cover letters for internships and jobs."
    )

    if st.button("Generate Cover Letter"):

        if uploaded_file and job_description:

            try:

                with st.spinner("Generating Cover Letter..."):

                    cover_prompt = f"""
                    You are a professional recruiter and career coach.

                    Generate a professional cover letter.

                    Requirements:
                    - Professional tone
                    - ATS-friendly wording
                    - Personalized using resume
                    - Personalized using job description
                    - Strong introduction
                    - Highlight technical skills
                    - Show enthusiasm
                    - End professionally

                    Resume:
                    {resume_text}

                    Job Description:
                    {job_description}
                    """

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[
                            {
                                "role": "system",
                                "content": "You are an expert cover letter writer."
                            },
                            {
                                "role": "user",
                                "content": cover_prompt
                            }
                        ]
                    )

                    cover_letter = response.choices[0].message.content

                st.subheader("📄 Generated Cover Letter")

                st.write(cover_letter)

                # SAVE DOCX
                doc = Document()

                doc.add_paragraph(cover_letter)

                doc.save("Cover_Letter.docx")

                # SAVE PDF
                pdf_file = "Cover_Letter.pdf"

                c = canvas.Canvas(pdf_file)

                textobject = c.beginText(40, 800)

                for line in cover_letter.split('\n'):
                    textobject.textLine(line)

                c.drawText(textobject)

                c.save()

                # DOWNLOAD DOCX
                with open("Cover_Letter.docx", "rb") as file:

                    st.download_button(
                        label="⬇ Download DOCX Cover Letter",
                        data=file,
                        file_name="Cover_Letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # DOWNLOAD PDF
                with open("Cover_Letter.pdf", "rb") as file:

                    st.download_button(
                        label="⬇ Download PDF Cover Letter",
                        data=file,
                        file_name="Cover_Letter.pdf",
                        mime="application/pdf"
                    )

                st.success(
                    "✅ Professional Cover Letter Generated Successfully"
                )

                speak(
                    f"{user_name}, your professional cover letter is ready."
                )

            except Exception as e:

                st.error(f"Error: {e}")

        else:

            st.error(
                "Please upload resume and paste job description."
            )